# -*- coding: utf-8 -*-
"""
DOCX → EPUB 변환기
- 리소스: resData/
- 결과물: saveResult/
- 이미지·표 포함, UTF-8, JS 코드 2칸 들여쓰기 보존
"""

from __future__ import annotations

import html
import io
import re
import sys
import uuid
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from ebooklib import epub
from PIL import Image

# ---------------------------------------------------------------------------
# 경로 / 상수
# ---------------------------------------------------------------------------
ROOT = Path(__file__).resolve().parent
RES_DIR = ROOT / "resData"
OUT_DIR = ROOT / "saveResult"

MONO_FONT_KEYWORDS = (
    "mono",
    "consolas",
    "courier",
    "cascadia",
    "source code",
    "fira code",
    "menlo",
    "monaco",
    "dejavu sans mono",
    "ubuntu mono",
    "lucida console",
)

HEADING_TAG = {
    "Heading 1": "h1",
    "Heading 2": "h2",
    "Heading 3": "h3",
    "Heading 4": "h4",
    "Heading 5": "h5",
    "Title": "h1",
    "Subtitle": "h2",
}

CSS = """\
@charset "UTF-8";
body {
  font-family: "Noto Sans KR", "Malgun Gothic", "Apple SD Gothic Neo", sans-serif;
  line-height: 1.7;
  color: #222;
  margin: 1em;
}
h1, h2, h3, h4, h5 {
  font-family: "Noto Sans KR", "Malgun Gothic", "Apple SD Gothic Neo", sans-serif;
  line-height: 1.35;
  margin: 1.4em 0 0.6em;
}
p {
  margin: 0.55em 0;
  word-break: keep-all;
  overflow-wrap: break-word;
}
pre.code, pre.code code {
  font-family: "Roboto Mono", Consolas, "Courier New", monospace;
  font-size: 0.82em;
  line-height: 1.45;
  white-space: pre;
  tab-size: 2;
  -moz-tab-size: 2;
}
pre.code {
  background: #f5f5f5;
  border: 1px solid #ddd;
  border-radius: 4px;
  padding: 0.75em 1em;
  overflow-x: auto;
  margin: 0.8em 0;
}
table {
  border-collapse: collapse;
  width: 100%;
  margin: 1em 0;
  font-size: 0.92em;
}
th, td {
  border: 1px solid #bbb;
  padding: 0.45em 0.6em;
  vertical-align: top;
  white-space: pre-wrap;
  font-family: "Roboto Mono", Consolas, "Courier New", monospace;
  word-break: break-word;
}
th {
  background: #eee;
  font-weight: bold;
}
img {
  max-width: 100%;
  height: auto;
  display: block;
  margin: 0.8em auto;
}
.caption {
  text-align: center;
  font-size: 0.9em;
  color: #555;
}
strong { font-weight: bold; }
em { font-style: italic; }
"""


def ensure_utf8_stdio() -> None:
    for stream in (sys.stdout, sys.stderr):
        if hasattr(stream, "reconfigure"):
            stream.reconfigure(encoding="utf-8")


def is_monospace_font(name: str | None) -> bool:
    if not name:
        return False
    lower = name.lower()
    return any(k in lower for k in MONO_FONT_KEYWORDS)


def run_font_names(run) -> set[str]:
    names: set[str] = set()
    if run.font.name:
        names.add(run.font.name)
    r_pr = run._element.find(qn("w:rPr"))
    if r_pr is not None:
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is not None:
            for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
                value = r_fonts.get(qn(f"w:{attr}"))
                if value:
                    names.add(value)
    return names


def has_monospace_run(paragraph: Paragraph) -> bool:
    style_name = (paragraph.style.name if paragraph.style else "") or ""
    if "code" in style_name.lower():
        return True
    for run in paragraph.runs:
        fonts = run_font_names(run)
        if any(is_monospace_font(f) for f in fonts):
            return True
    return False


def is_code_paragraph(paragraph: Paragraph, in_code_block: bool = False) -> bool:
    """
    코드 단락 판별.
    - Roboto Mono 등 모노스페이스 → 코드
    - 코드 블록 진행 중이면 선행 들여쓰기(스페이스/탭)가 있는 줄도 코드로 포함
      (한글·특수문자(❶❷)만 Arial로 들어간 줄 대응)
    """
    if has_monospace_run(paragraph):
        return True

    if not in_code_block:
        return False

    text = paragraph_raw_text(paragraph)
    if text.startswith(" ") or text.startswith("\t"):
        return True

    stripped = text.strip()
    if not stripped:
        return False

    # 들여쓰기 없는 닫는 토큰 등 (폰트 정보가 빠진 경우 대비)
    if stripped.startswith(("</", "});", "})", "};", "}")):
        return True
    return False


def paragraph_raw_text(paragraph: Paragraph) -> str:
    """들여쓰기(선행 공백)를 포함한 단락 텍스트. 소프트 줄바꿈은 \\n으로."""
    parts: list[str] = []
    for child in paragraph._element.iter():
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "t":
            parts.append(child.text or "")
        elif tag == "tab":
            parts.append("  ")  # 탭 → 스페이스 2칸
        elif tag == "br":
            parts.append("\n")
        elif tag == "cr":
            parts.append("\n")
    return "".join(parts)


def is_blank_paragraph(paragraph: Paragraph) -> bool:
    if paragraph_raw_text(paragraph).strip():
        return False
    # 이미지만 있는 단락은 blank가 아님
    if paragraph._element.xpath(".//a:blip"):
        return False
    return True


def escape_text(text: str) -> str:
    return html.escape(text, quote=True)


def format_inline_runs(paragraph: Paragraph) -> str:
    """일반 단락의 인라인(굵게/기울임) HTML."""
    chunks: list[str] = []
    for run in paragraph.runs:
        text = run.text or ""
        if not text:
            # 줄바꿈 런
            if run._element.xpath(".//w:br") or run._element.xpath(".//w:cr"):
                chunks.append("<br/>")
            continue
        escaped = escape_text(text).replace("\n", "<br/>")
        if run.bold:
            escaped = f"<strong>{escaped}</strong>"
        if run.italic:
            escaped = f"<em>{escaped}</em>"
        chunks.append(escaped)

    if chunks:
        return "".join(chunks)

    # runs가 비어 있으면 XML 텍스트 fallback
    return escape_text(paragraph_raw_text(paragraph)).replace("\n", "<br/>")


def normalize_js_indent(lines: list[str]) -> list[str]:
    """
    코드 블록의 들여쓰기를 스페이스 2칸 단위로 정규화.
    - 탭 → 스페이스 2칸
    - 선행 공백 개수를 가장 가까운 짝수(2의 배수)로 맞춤
    - 원본 상대 들여쓰기 구조는 유지
    """
    normalized: list[str] = []
    for line in lines:
        # 탭을 스페이스 2칸으로
        expanded = line.replace("\t", "  ")
        # 선행 공백 추출
        stripped = expanded.lstrip(" ")
        if not stripped:
            normalized.append("")
            continue
        lead = len(expanded) - len(stripped)
        # 홀수 스페이스가 있으면 가장 가까운 짝수로 (내림이 아니라 반올림에 가깝게)
        # 예: 1→2, 3→4 … 다만 0은 유지. 문서가 이미 2칸이면 그대로.
        if lead % 2 == 1:
            lead += 1
        normalized.append((" " * lead) + stripped)
    return normalized


def iter_block_items(document: Document):
    """본문 순서대로 Paragraph / Table yield."""
    parent = document.element.body
    for child in parent.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            yield Paragraph(child, document)
        elif tag == "tbl":
            yield Table(child, document)


def image_media_type(ext: str) -> str:
    ext = ext.lower().lstrip(".")
    mapping = {
        "png": "image/png",
        "jpg": "image/jpeg",
        "jpeg": "image/jpeg",
        "gif": "image/gif",
        "svg": "image/svg+xml",
        "webp": "image/webp",
        "bmp": "image/bmp",
        "tif": "image/tiff",
        "tiff": "image/tiff",
    }
    return mapping.get(ext, "image/png")


class EpubBuilder:
    def __init__(self, title: str, author: str = "Unknown"):
        self.book = epub.EpubBook()
        self.book.set_identifier(str(uuid.uuid4()))
        self.book.set_title(title)
        self.book.set_language("ko")
        self.book.add_author(author)

        self.style = epub.EpubItem(
            uid="style",
            file_name="style/default.css",
            media_type="text/css",
            content=CSS.encode("utf-8"),
        )
        self.book.add_item(self.style)

        self.chapters: list[epub.EpubHtml] = []
        self._chapter_html: list[str] = []
        self._chapter_title = title
        self._chapter_index = 0
        self._image_index = 0
        self._image_cache: dict[str, str] = {}  # rel_id → epub file name
        self._spine_images: list[epub.EpubItem] = []

    def flush_chapter(self) -> None:
        if not self._chapter_html:
            return
        self._chapter_index += 1
        file_name = f"chap_{self._chapter_index:03d}.xhtml"
        chapter = epub.EpubHtml(
            title=self._chapter_title or f"Chapter {self._chapter_index}",
            file_name=file_name,
            lang="ko",
        )
        # ebooklib이 XHTML 래퍼를 붙이므로 body 조각만 전달 (UTF-8)
        chapter.content = "\n".join(self._chapter_html)
        chapter.add_item(self.style)
        self.book.add_item(chapter)
        self.chapters.append(chapter)
        self._chapter_html = []

    def start_chapter(self, title: str) -> None:
        self.flush_chapter()
        self._chapter_title = title.strip() or f"Section {self._chapter_index + 1}"

    def append(self, html_fragment: str) -> None:
        self._chapter_html.append(html_fragment)

    def add_image_blob(self, blob: bytes, preferred_name: str, rel_key: str) -> str | None:
        """이미지를 EPUB에 추가하고 상대 경로 반환. 너무 작은 스페이서는 생략."""
        if rel_key in self._image_cache:
            return self._image_cache[rel_key]

        if len(blob) < 200:
            return None

        ext = Path(preferred_name).suffix.lower().lstrip(".") or "png"
        # BMP 등은 PNG로 재인코딩해 호환성 확보
        try:
            with Image.open(io.BytesIO(blob)) as im:
                if ext in ("bmp", "tif", "tiff") or im.format not in (
                    "PNG",
                    "JPEG",
                    "GIF",
                    "WEBP",
                ):
                    buf = io.BytesIO()
                    if im.mode not in ("RGB", "RGBA", "L", "P"):
                        im = im.convert("RGBA")
                    im.save(buf, format="PNG")
                    blob = buf.getvalue()
                    ext = "png"
                else:
                    # 원본 유지하되 확장자 정리
                    fmt = (im.format or "PNG").lower()
                    if fmt == "jpeg":
                        ext = "jpg"
                    elif fmt == "png":
                        ext = "png"
                    elif fmt == "gif":
                        ext = "gif"
                    elif fmt == "webp":
                        ext = "webp"
        except Exception:
            # PIL 실패 시 원본 blob 그대로 사용
            pass

        self._image_index += 1
        file_name = f"images/img_{self._image_index:03d}.{ext}"
        item = epub.EpubItem(
            uid=f"img_{self._image_index:03d}",
            file_name=file_name,
            media_type=image_media_type(ext),
            content=blob,
        )
        self.book.add_item(item)
        self._spine_images.append(item)
        self._image_cache[rel_key] = file_name
        return file_name

    def finish(self, out_path: Path) -> None:
        if not self._chapter_html and not self.chapters:
            self.append("<p>(내용 없음)</p>")
        self.flush_chapter()

        self.book.toc = tuple(
            epub.Link(ch.file_name, ch.title, ch.get_id()) for ch in self.chapters
        )
        self.book.add_item(epub.EpubNcx())
        self.book.add_item(epub.EpubNav())
        self.book.spine = ["nav"] + self.chapters

        out_path.parent.mkdir(parents=True, exist_ok=True)
        epub.write_epub(str(out_path), self.book, {"epub3_titlepage": False})


def extract_images_html(paragraph: Paragraph, document: Document, builder: EpubBuilder) -> str:
    parts: list[str] = []
    for blip in paragraph._element.xpath(".//a:blip"):
        embed = blip.get(qn("r:embed"))
        if not embed or embed not in document.part.rels:
            continue
        rel = document.part.rels[embed]
        if "image" not in rel.reltype:
            continue
        blob = rel.target_part.blob
        target = rel.target_ref
        path = builder.add_image_blob(blob, target, embed)
        if path:
            parts.append(f'<img src="{path}" alt="image"/>')
    return "\n".join(parts)


def table_to_html(table: Table) -> str:
    rows_html: list[str] = []
    for r_idx, row in enumerate(table.rows):
        cells_html: list[str] = []
        for cell in row.cells:
            # 셀 안 여러 단락 → 줄바꿈 유지 (코드/도식용)
            cell_lines: list[str] = []
            for p in cell.paragraphs:
                raw = paragraph_raw_text(p)
                # 탭 → 2스페이스, 들여쓰기 정규화는 줄 단위
                raw = raw.replace("\t", "  ")
                cell_lines.append(raw)
            cell_text = "\n".join(cell_lines)
            # 들여쓰기 정규화
            cell_text = "\n".join(normalize_js_indent(cell_text.split("\n")))
            tag = "th" if r_idx == 0 else "td"
            cells_html.append(f"<{tag}>{escape_text(cell_text)}</{tag}>")
        rows_html.append("<tr>" + "".join(cells_html) + "</tr>")
    return "<table>\n" + "\n".join(rows_html) + "\n</table>"


def convert_docx_to_epub(docx_path: Path, out_path: Path) -> None:
    document = Document(str(docx_path))

    # 제목: Heading 1 우선, 없으면 파일명
    title = docx_path.stem
    for p in document.paragraphs[:30]:
        if p.style and p.style.name == "Heading 1" and p.text.strip():
            title = re.sub(r"\s+", " ", p.text).strip()
            break

    builder = EpubBuilder(title=title)

    code_buffer: list[str] = []
    started = False

    def flush_code() -> None:
        nonlocal code_buffer
        if not code_buffer:
            return
        # 끝의 빈 줄 제거 (본문과의 구분용 blank)
        while code_buffer and code_buffer[-1] == "":
            code_buffer.pop()
        lines = normalize_js_indent(code_buffer)
        # 줄마다 escape 후 \n으로 join — pre가 공백·개행 보존
        code_html = escape_text("\n".join(lines))
        builder.append(f'<pre class="code"><code>{code_html}</code></pre>')
        code_buffer = []

    for block in iter_block_items(document):
        if isinstance(block, Table):
            flush_code()
            builder.append(table_to_html(block))
            started = True
            continue

        paragraph: Paragraph = block
        style_name = paragraph.style.name if paragraph.style else "normal"
        images_html = extract_images_html(paragraph, document, builder)

        # 코드 단락 (진행 중이면 들여쓰기 줄도 코드로 흡수)
        if is_code_paragraph(paragraph, in_code_block=bool(code_buffer)):
            raw = paragraph_raw_text(paragraph)
            # 단락 내부 soft-break 는 여러 줄일 수 있음
            for line in raw.split("\n"):
                code_buffer.append(line)
            if images_html:
                flush_code()
                builder.append(images_html)
            started = True
            continue

        # 코드 블록 중간의 빈 줄은 코드에 포함
        if code_buffer and is_blank_paragraph(paragraph) and not images_html:
            code_buffer.append("")
            continue

        flush_code()

        if images_html:
            builder.append(images_html)

        text = paragraph_raw_text(paragraph)
        if not text.strip() and not images_html:
            continue

        heading_tag = HEADING_TAG.get(style_name)
        if heading_tag:
            # 챕터 분기: h1/h2/h3
            clean_title = re.sub(r"\s+", " ", text).strip()
            if heading_tag in ("h1", "h2", "h3"):
                if started or builder.chapters or builder._chapter_html:
                    builder.start_chapter(clean_title)
                else:
                    builder._chapter_title = clean_title
            inline = format_inline_runs(paragraph)
            builder.append(f"<{heading_tag}>{inline}</{heading_tag}>")
            started = True
            continue

        inline = format_inline_runs(paragraph)
        if inline.strip():
            builder.append(f"<p>{inline}</p>")
            started = True

    flush_code()
    builder.finish(out_path)
    print(f"[OK] {docx_path.name} → {out_path}")


def main() -> int:
    ensure_utf8_stdio()

    # 가상환경 확인
    in_venv = sys.prefix != getattr(sys, "base_prefix", sys.prefix)
    venv_hint = Path(sys.prefix).name
    if not in_venv or venv_hint != ".epubvenv":
        print(
            "[경고] .epubvenv 가상환경 밖에서 실행 중인 것으로 보입니다.\n"
            f"       현재 prefix: {sys.prefix}\n"
            "       .\\.epubvenv\\Scripts\\python.exe docx_to_epub.py 로 실행하세요.",
            file=sys.stderr,
        )

    if not RES_DIR.is_dir():
        print(f"[오류] 리소스 폴더가 없습니다: {RES_DIR}", file=sys.stderr)
        return 1

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    docx_files = sorted(RES_DIR.glob("*.docx"))
    if not docx_files:
        print(f"[오류] {RES_DIR} 에 docx 파일이 없습니다.", file=sys.stderr)
        return 1

    for docx_path in docx_files:
        if docx_path.name.startswith("~$"):
            continue
        out_path = OUT_DIR / (docx_path.stem + ".epub")
        convert_docx_to_epub(docx_path, out_path)

    print(f"완료: {len(docx_files)}개 파일 변환 → {OUT_DIR}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
